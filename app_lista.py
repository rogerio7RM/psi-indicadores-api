# app_lista.py
# Example env vars for email CLI (powershell):
# $env:REPORT_SMTP_HOST = "smtp.gmail.com"
# $env:REPORT_SMTP_USER = "rogmoues@gmail.com"
# $env:REPORT_SMTP_PASSWORD = "sglq chbt cmps dbei"
# $env:REPORT_EMAIL_TO = "rogerio7@gmail.com"
# python app_lista.py --send-email --tickers "QQQ,SPY"

#pARA COLOCAR NO TASK MANAGER:
#-NoProfile -ExecutionPolicy Bypass -File "C:\PrimeSphere\Indicadores\send_report.ps1"

#Geral: dÃª um nome (ex.: â€œPrimeSphere Reportâ€), selecione â€œExecutar estando o usuÃ¡rio conectado ou nÃ£oâ€ e marque â€œExecutar com privilÃ©gios mais altosâ€ se a conta precisar.
#Triggers: â€œNovoâ€¦â€ â†’ Iniciar tarefa â€œDiariamenteâ€ â†’ escolha horÃ¡rio â†’ OK.
#AÃ§Ãµes: â€œNovoâ€¦â€ â†’ AÃ§Ã£o â€œIniciar um programaâ€ â†’ Programa powershell.exe â†’ argumento:
#-NoProfile -ExecutionPolicy Bypass -File "C:\PrimeSphere\Indicadores\send_report.ps1"
#CondiÃ§Ãµes/ConfiguraÃ§Ãµes: desmarque â€œIniciar somente se estiver usando energia CAâ€ caso queira rodar em notebooks, e habilite â€œExecutar o mais rÃ¡pido possÃ­vel apÃ³s uma inicializaÃ§Ã£o agendada perdidaâ€.



import os
import warnings
import math
import smtplib
import ssl
from email.message import EmailMessage
from email.utils import formatdate
from flask import Flask, render_template, request, send_file, abort, jsonify
import yfinance as yf
import pandas as pd
import ta
from datetime import datetime
import io
from functools import lru_cache
try:
    from PIL import Image
except Exception:
    Image = None
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Base URL for the charting app (app.py)
app = Flask(__name__)

GRAPH_APP_URL = os.getenv("GRAPH_APP_URL", "http://localhost:5000")
DEFAULT_TICKERS = "QQQ,GLD, SLV, GOOGL, SPY,LLY, PLTR,AMD,BTC,AMZN,SOFI,TSLA,NVDA,NFLX"

SUPPORTED_LANGS = {"pt", "en", "es"}
DEFAULT_LANG = "pt"
I18N = {
    "pt": {
        "language": "Idioma",
        "language_pt": "PT",
        "language_en": "EN",
        "language_es": "ES",
        "list_title": "Lista de Indicadores & Recomendacoes",
        "list_subtitle": "Analise varios tickers e abra graficos detalhados com um clique.",
        "list_form_label": "Tickers (separados por virgula)",
        "list_form_placeholder": "ex.: QQQ, COWZ, IBIT",
        "list_analyze": "Analisar lista",
        "list_hint": "Clique no ticker para abrir o grafico em nova aba.",
        "list_dcf_button": "DCF",
        "list_results": "Resultados",
        "download_report": "Baixar relatorio Word",
        "list_table_ticker": "Ticker",
        "list_table_price": "Preco",
        "list_table_macd": "MACD / Sinal",
        "list_table_recommendation": "Recomendacao",
        "list_table_summary": "Resumo / Racional",
        "list_empty": "Digite um ou mais tickers acima e clique em \"Analisar lista\".",
        "unavailable": "Indisponivel",
        "error": "Erro",
        "price_live": "ao vivo",
        "price_prev_close": "fechamento ant",
        "change_today": "Hoje",
        "change_prev_close": "Fech. ant",
        "dcf_title": "DCF estruturado",
        "dcf_subtitle": "Modelo quantitativo com dados reais mais recentes da companhia.",
        "dcf_latest_title": "Dados reais recentes",
        "dcf_history_title": "Historico recente (anual)",
        "dcf_assumptions_title": "Premissas",
        "dcf_model_title": "Modelo DCF",
        "dcf_structured_title": "DCF estruturado (passo a passo)",
        "dcf_base_title": "Dados base",
        "dcf_projection_title": "Projecao de FCF",
        "dcf_terminal_title": "Valor terminal",
        "dcf_discount_title": "Desconto a valor presente",
        "dcf_ev_title": "Enterprise value estimado",
        "dcf_interpretation_title": "Interpretacao",
        "dcf_sensitivity_summary_title": "Sensibilidade (cenarios)",
        "dcf_conclusion_title": "Conclusao",
        "dcf_pv_sum": "Soma PV (anos 1-{years})",
        "dcf_market_cap": "Market cap",
        "dcf_sensitivity_title": "Sensibilidade",
        "dcf_notes_title": "Notas",
        "dcf_field_period": "Periodo",
        "dcf_field_revenue": "Receita",
        "dcf_field_operating_cf": "Caixa operacional",
        "dcf_field_capex": "Capex",
        "dcf_field_fcf": "Free cash flow",
        "dcf_field_ebitda": "EBITDA",
        "dcf_field_net_income": "Lucro liquido",
        "dcf_field_cash": "Caixa",
        "dcf_field_debt": "Divida total",
        "dcf_field_shares": "Acoes em circulacao",
        "dcf_field_fiscal": "Ultimo ano fiscal",
        "dcf_field_currency": "Moeda",
        "dcf_field_years": "Horizonte (anos)",
        "dcf_field_discount_rate": "Taxa de desconto",
        "dcf_field_terminal_growth": "Crescimento terminal",
        "dcf_field_growth": "Crescimento FCF usado",
        "dcf_field_growth_source": "Fonte do crescimento",
        "dcf_table_year": "Ano",
        "dcf_table_fcf": "FCF proj.",
        "dcf_table_discount": "Fator desc.",
        "dcf_table_pv": "Valor presente",
        "dcf_terminal_value": "Valor terminal",
        "dcf_terminal_pv": "Valor terminal (VP)",
        "dcf_summary_enterprise": "Valor da firma (EV)",
        "dcf_summary_equity": "Valor do equity",
        "dcf_summary_intrinsic": "Valor intrinseco por acao",
        "dcf_summary_price": "Preco de mercado",
        "dcf_summary_upside": "Upside",
        "dcf_per_share_note": "Valores por acao",
        "dcf_missing_data": "Dados insuficientes para calcular o DCF.",
        "dcf_note_source": "Fonte: Yahoo Finance (yfinance).",
        "dcf_note_simplified": "Modelo DCF simplificado para referencia, nao e recomendacao de investimento.",
        "dcf_note_terminal_adjust": "Crescimento terminal ajustado para ficar abaixo da taxa de desconto.",
    },
    "en": {
        "language": "Language",
        "language_pt": "PT",
        "language_en": "EN",
        "language_es": "ES",
        "list_title": "Indicator List & Recommendations",
        "list_subtitle": "Analyze multiple tickers at once and open detailed charts with a single click.",
        "list_form_label": "Tickers (comma separated)",
        "list_form_placeholder": "e.g., QQQ, COWZ, IBIT",
        "list_analyze": "Analyze list",
        "list_hint": "Click the ticker to open the charting app in a new tab.",
        "list_dcf_button": "DCF",
        "list_results": "Results",
        "download_report": "Download Word report",
        "list_table_ticker": "Ticker",
        "list_table_price": "Price",
        "list_table_macd": "MACD / Signal",
        "list_table_recommendation": "Recommendation",
        "list_table_summary": "Summary / Rationale",
        "list_empty": "Enter one or more tickers above and click \"Analyze list\".",
        "unavailable": "Unavailable",
        "error": "Error",
        "price_live": "live",
        "price_prev_close": "prev close",
        "change_today": "Today",
        "change_prev_close": "Prev close",
        "dcf_title": "Structured DCF",
        "dcf_subtitle": "Quantitative model using the latest real company data.",
        "dcf_latest_title": "Latest real data",
        "dcf_history_title": "Recent history (annual)",
        "dcf_assumptions_title": "Assumptions",
        "dcf_model_title": "DCF model",
        "dcf_structured_title": "Structured DCF (step by step)",
        "dcf_base_title": "Base data",
        "dcf_projection_title": "FCF projection",
        "dcf_terminal_title": "Terminal value",
        "dcf_discount_title": "Discount to present value",
        "dcf_ev_title": "Estimated enterprise value",
        "dcf_interpretation_title": "Interpretation",
        "dcf_sensitivity_summary_title": "Sensitivity (scenarios)",
        "dcf_conclusion_title": "Conclusion",
        "dcf_pv_sum": "PV sum (years 1-{years})",
        "dcf_market_cap": "Market cap",
        "dcf_sensitivity_title": "Sensitivity",
        "dcf_notes_title": "Notes",
        "dcf_field_period": "Period",
        "dcf_field_revenue": "Revenue",
        "dcf_field_operating_cf": "Operating cash flow",
        "dcf_field_capex": "Capex",
        "dcf_field_fcf": "Free cash flow",
        "dcf_field_ebitda": "EBITDA",
        "dcf_field_net_income": "Net income",
        "dcf_field_cash": "Cash",
        "dcf_field_debt": "Total debt",
        "dcf_field_shares": "Shares outstanding",
        "dcf_field_fiscal": "Latest fiscal year",
        "dcf_field_currency": "Currency",
        "dcf_field_years": "Horizon (years)",
        "dcf_field_discount_rate": "Discount rate",
        "dcf_field_terminal_growth": "Terminal growth",
        "dcf_field_growth": "FCF growth used",
        "dcf_field_growth_source": "Growth source",
        "dcf_table_year": "Year",
        "dcf_table_fcf": "Projected FCF",
        "dcf_table_discount": "Discount factor",
        "dcf_table_pv": "Present value",
        "dcf_terminal_value": "Terminal value",
        "dcf_terminal_pv": "Terminal value (PV)",
        "dcf_summary_enterprise": "Enterprise value (EV)",
        "dcf_summary_equity": "Equity value",
        "dcf_summary_intrinsic": "Intrinsic value per share",
        "dcf_summary_price": "Market price",
        "dcf_summary_upside": "Upside",
        "dcf_per_share_note": "Per-share values",
        "dcf_missing_data": "Insufficient data to calculate the DCF.",
        "dcf_note_source": "Source: Yahoo Finance (yfinance).",
        "dcf_note_simplified": "Simplified DCF model for reference, not investment advice.",
        "dcf_note_terminal_adjust": "Terminal growth adjusted to stay below the discount rate.",
    },
    "es": {
        "language": "Idioma",
        "language_pt": "PT",
        "language_en": "EN",
        "language_es": "ES",
        "list_title": "Lista de Indicadores y Recomendaciones",
        "list_subtitle": "Analiza varios tickers y abre graficos detallados con un clic.",
        "list_form_label": "Tickers (separados por coma)",
        "list_form_placeholder": "ej.: QQQ, COWZ, IBIT",
        "list_analyze": "Analizar lista",
        "list_hint": "Haz clic en el ticker para abrir el grafico en una nueva pestana.",
        "list_dcf_button": "DCF",
        "list_results": "Resultados",
        "download_report": "Descargar reporte Word",
        "list_table_ticker": "Ticker",
        "list_table_price": "Precio",
        "list_table_macd": "MACD / Senal",
        "list_table_recommendation": "Recomendacion",
        "list_table_summary": "Resumen / Razonamiento",
        "list_empty": "Ingresa uno o mas tickers arriba y haz clic en \"Analizar lista\".",
        "unavailable": "No disponible",
        "error": "Error",
        "price_live": "en vivo",
        "price_prev_close": "cierre prev",
        "change_today": "Hoy",
        "change_prev_close": "Cierre prev",
        "dcf_title": "DCF estructurado",
        "dcf_subtitle": "Modelo cuantitativo con datos reales mas recientes de la compania.",
        "dcf_latest_title": "Datos reales recientes",
        "dcf_history_title": "Historia reciente (anual)",
        "dcf_assumptions_title": "Supuestos",
        "dcf_model_title": "Modelo DCF",
        "dcf_structured_title": "DCF estructurado (paso a paso)",
        "dcf_base_title": "Datos base",
        "dcf_projection_title": "Proyeccion de FCF",
        "dcf_terminal_title": "Valor terminal",
        "dcf_discount_title": "Descuento a valor presente",
        "dcf_ev_title": "Enterprise value estimado",
        "dcf_interpretation_title": "Interpretacion",
        "dcf_sensitivity_summary_title": "Sensibilidad (escenarios)",
        "dcf_conclusion_title": "Conclusion",
        "dcf_pv_sum": "Suma PV (anos 1-{years})",
        "dcf_market_cap": "Market cap",
        "dcf_sensitivity_title": "Sensibilidad",
        "dcf_notes_title": "Notas",
        "dcf_field_period": "Periodo",
        "dcf_field_revenue": "Ingresos",
        "dcf_field_operating_cf": "Flujo operativo",
        "dcf_field_capex": "Capex",
        "dcf_field_fcf": "Free cash flow",
        "dcf_field_ebitda": "EBITDA",
        "dcf_field_net_income": "Utilidad neta",
        "dcf_field_cash": "Caja",
        "dcf_field_debt": "Deuda total",
        "dcf_field_shares": "Acciones en circulacion",
        "dcf_field_fiscal": "Ultimo ano fiscal",
        "dcf_field_currency": "Moneda",
        "dcf_field_years": "Horizonte (anos)",
        "dcf_field_discount_rate": "Tasa de descuento",
        "dcf_field_terminal_growth": "Crecimiento terminal",
        "dcf_field_growth": "Crecimiento FCF usado",
        "dcf_field_growth_source": "Fuente de crecimiento",
        "dcf_table_year": "Ano",
        "dcf_table_fcf": "FCF proyectado",
        "dcf_table_discount": "Factor desc.",
        "dcf_table_pv": "Valor presente",
        "dcf_terminal_value": "Valor terminal",
        "dcf_terminal_pv": "Valor terminal (VP)",
        "dcf_summary_enterprise": "Valor de la firma (EV)",
        "dcf_summary_equity": "Valor del equity",
        "dcf_summary_intrinsic": "Valor intrinseco por accion",
        "dcf_summary_price": "Precio de mercado",
        "dcf_summary_upside": "Upside",
        "dcf_per_share_note": "Valores por accion",
        "dcf_missing_data": "Datos insuficientes para calcular el DCF.",
        "dcf_note_source": "Fuente: Yahoo Finance (yfinance).",
        "dcf_note_simplified": "Modelo DCF simplificado para referencia, no es recomendacion de inversion.",
        "dcf_note_terminal_adjust": "Crecimiento terminal ajustado para quedar por debajo de la tasa de descuento.",
    },
}


def resolve_lang(raw):
    lang = (raw or "").strip().lower()
    return lang if lang in SUPPORTED_LANGS else DEFAULT_LANG


ANALYSIS_TEXT = {
    "pt": {
        "insights_unavailable": "Indicadores indisponiveis: dados insuficientes.",
        "rec_entry": "ENTRAR",
        "rec_exit": "SAIR",
        "rec_neutral": "NEUTRO",
        "trend_positive": "Tendencia positiva (preco > EMA9 > EMA26)",
        "macd_hist_pos": "Histograma MACD > 0",
        "rsi_between": "RSI entre {min}-{max}",
        "adx_confirms": "ADX confirma a estrutura",
        "trend_strength": "Forca da tendencia (ADX >= 25)",
        "trend_weakening": "Tendencia enfraquecendo ou MACD cruzando para baixo",
        "rsi_take_profit": "RSI >= {max} (realizar lucro)",
        "overbought_pullback": "Sobrecomprado e rompendo a banda superior: espere pullback",
        "oversold_reversal": "Sobrevendido com rompimento da banda inferior: aguarde reversao",
        "trend_up_wait": "Tendencia de alta; aguarde pullback",
        "trend_down_wait": "Tendencia de baixa; aguarde confirmacao de reversao",
        "sideways_wait": "Estrutura lateral; aguarde um gatilho",
        "summary_entry": "Tendencia e momentum alinham para cima.",
        "summary_exit": "Momentum enfraquecendo; proteja capital.",
        "summary_overbought": "Sobrecomprado com rompimento de banda; espere pullback.",
        "summary_oversold": "Sobrevendido com rompimento de banda; observe reversao.",
        "summary_mixed": "Estrutura mista; aguarde confirmacao.",
        "summary_none": "Sem sinais fortes; mantenha paciencia.",
    },
    "en": {
        "insights_unavailable": "Indicators unavailable: insufficient valid data.",
        "rec_entry": "ENTRY (buy)",
        "rec_exit": "EXIT / AVOID",
        "rec_neutral": "NEUTRAL / WAIT",
        "trend_positive": "Trend positive (price > EMA9 > EMA26)",
        "macd_hist_pos": "MACD histogram > 0",
        "rsi_between": "RSI between {min}-{max}",
        "adx_confirms": "ADX confirms structure",
        "trend_strength": "Trend strength (ADX >= 25)",
        "trend_weakening": "Trend weakening or MACD crossing down",
        "rsi_take_profit": "RSI >= {max} (take profit)",
        "overbought_pullback": "Overbought and breaking the upper band: expect a pullback",
        "oversold_reversal": "Oversold with a lower band break: wait for a reversal",
        "trend_up_wait": "Trend up; wait for a pullback signal",
        "trend_down_wait": "Trend down; wait for reversal confirmation",
        "sideways_wait": "Sideways structure; wait for a trigger",
        "summary_entry": "Trend and momentum align to the upside.",
        "summary_exit": "Momentum fading; protect capital.",
        "summary_overbought": "Overbought with a band break; expect a pullback.",
        "summary_oversold": "Oversold with a band break; watch for reversal.",
        "summary_mixed": "Mixed structure; wait for confirmation.",
        "summary_none": "No strong signals detected; stay patient.",
    },
    "es": {
        "insights_unavailable": "Indicadores no disponibles: datos insuficientes.",
        "rec_entry": "ENTRADA (compra)",
        "rec_exit": "SALIR / EVITAR",
        "rec_neutral": "NEUTRO / ESPERAR",
        "trend_positive": "Tendencia positiva (precio > EMA9 > EMA26)",
        "macd_hist_pos": "Histograma MACD > 0",
        "rsi_between": "RSI entre {min}-{max}",
        "adx_confirms": "ADX confirma la estructura",
        "trend_strength": "Fuerza de la tendencia (ADX >= 25)",
        "trend_weakening": "Tendencia debilitandose o MACD cruzando a la baja",
        "rsi_take_profit": "RSI >= {max} (tomar ganancia)",
        "overbought_pullback": "Sobrecomprado y rompiendo la banda superior: espere retroceso",
        "oversold_reversal": "Sobrevendido con ruptura de la banda inferior: esperar rebote",
        "trend_up_wait": "Tendencia alcista; esperar retroceso",
        "trend_down_wait": "Tendencia bajista; esperar confirmacion de giro",
        "sideways_wait": "Estructura lateral; esperar un disparador",
        "summary_entry": "Tendencia y momentum alineados al alza.",
        "summary_exit": "Momentum debilitandose; proteja capital.",
        "summary_overbought": "Sobrecomprado con ruptura de banda; espere retroceso.",
        "summary_oversold": "Sobrevendido con ruptura de banda; vigile rebote.",
        "summary_mixed": "Estructura mixta; espere confirmacion.",
        "summary_none": "Sin senales fuertes; sea paciente.",
    },
}

def get_logo_path():
    logo_candidates = [
        r"C:\Users\roger\OneDrive\PrimeSphere\Logo\PrimeShpere Logo.jpg",
        os.path.join(app.root_path, "static", "PrimeShpere Logo.jpg"),
    ]
    return next((candidate for candidate in logo_candidates if os.path.exists(candidate)), None)


def get_logo_stream_no_background(logo_path):
    if not logo_path or Image is None:
        return None
    try:
        mtime = os.path.getmtime(logo_path)
        png_bytes = _logo_png_bytes_cached(logo_path, mtime)
        if not png_bytes:
            return None
        stream = io.BytesIO(png_bytes)
        stream.seek(0)
        return stream
    except Exception:
        return None


@lru_cache(maxsize=4)
def _logo_png_bytes_cached(logo_path, mtime):
    if Image is None:
        return None
    try:
        with Image.open(logo_path).convert("RGBA") as img:
            cleaned = [
                (r, g, b, 0) if (r >= 245 and g >= 245 and b >= 245) else (r, g, b, a)
                for r, g, b, a in img.getdata()
            ]
            img.putdata(cleaned)
            stream = io.BytesIO()
            img.save(stream, format="PNG")
            return stream.getvalue()
    except Exception:
        return None

LOGO_PATH = get_logo_path()

DEFAULT_SMTP_HOST = "smtp.gmail.com"
DEFAULT_SMTP_PORT = 587
DEFAULT_SMTP_USER = "rogmoues@gmail.com"
DEFAULT_SMTP_PASSWORD = "Enribia1975"
DEFAULT_EMAIL_FROM = DEFAULT_SMTP_USER
DEFAULT_EMAIL_TO = "rogerio7@gmail.com"

REPORT_SMTP_HOST = os.getenv("REPORT_SMTP_HOST", DEFAULT_SMTP_HOST)
REPORT_SMTP_PORT = int(os.getenv("REPORT_SMTP_PORT", str(DEFAULT_SMTP_PORT)))
REPORT_SMTP_USER = os.getenv("REPORT_SMTP_USER", DEFAULT_SMTP_USER)
REPORT_SMTP_PASSWORD = os.getenv("REPORT_SMTP_PASSWORD", DEFAULT_SMTP_PASSWORD)
REPORT_SMTP_STARTTLS = os.getenv("REPORT_SMTP_STARTTLS", "true").lower() in ("1", "true", "yes", "on")
REPORT_EMAIL_FROM = os.getenv("REPORT_EMAIL_FROM") or REPORT_SMTP_USER or DEFAULT_EMAIL_FROM
REPORT_EMAIL_TO = os.getenv("REPORT_EMAIL_TO", DEFAULT_EMAIL_TO)
REPORT_EMAIL_SUBJECT = os.getenv("REPORT_EMAIL_SUBJECT", "PrimeSphere Daily Indicator Report")
REPORT_EMAIL_BODY = os.getenv("REPORT_EMAIL_BODY", "Segue o relatÃ³rio diÃ¡rio de indicadores da PrimeSphere.")
REPORT_EMAIL_BODY_HTML = os.getenv("REPORT_EMAIL_BODY_HTML")
REPORT_EMAIL_TICKERS = os.getenv("REPORT_EMAIL_TICKERS", DEFAULT_TICKERS)


warnings.filterwarnings(
    "ignore",
    message="invalid value encountered in scalar divide",
    category=RuntimeWarning,
    module=r"ta\.trend"
)
warnings.filterwarnings(
    "ignore",
    message="invalid value encountered in divide",
    category=RuntimeWarning,
    module=r"ta\.trend"
)


def normalize_ticker_string(raw):
    return (raw or "").upper().replace(" ", "")


def split_tickers(raw):
    return [t for t in normalize_ticker_string(raw).split(",") if t]


# =============================
# Utility functions
# =============================
RISK_OPT_PARAMS = {
    "ema_fast": 9,
    "ema_slow": 26,
    "rsi_entry_min": 50,
    "rsi_entry_max": 65,
    "rsi_exit_max": 70,
    "adx_min": 20,
    "atr_mult": 2.5,
    "use_cmf": False,
}

def _calc_prev_close_change_pct(df):
    if df is None or df.empty or "Close" not in df.columns:
        return None
    closes = df["Close"].dropna()
    if len(closes) < 3:
        return None
    prev_close = _safe_float(closes.iloc[-2])
    prev_prev_close = _safe_float(closes.iloc[-3])
    if prev_close is None or prev_prev_close in (None, 0):
        return None
    return ((prev_close - prev_prev_close) / prev_prev_close) * 100

def _safe_float(value):
    if value is None or pd.isna(value):
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _pick_float(mapping, keys):
    if not mapping:
        return None
    for key in keys:
        try:
            value = mapping.get(key)
        except Exception:
            value = None
        value = _safe_float(value)
        if value is not None:
            return value
    return None


def _ordered_statement_columns(df: pd.DataFrame):
    if df is None or df.empty:
        return []
    cols = list(df.columns)
    dated = []
    for col in cols:
        if isinstance(col, (pd.Timestamp, datetime)):
            dt = pd.Timestamp(col)
        else:
            try:
                dt = pd.to_datetime(col, errors="coerce")
            except Exception:
                dt = pd.NaT
        dated.append((col, dt))
    valid = [item for item in dated if pd.notna(item[1])]
    if valid:
        valid.sort(key=lambda x: x[1])
        return [c for c, _ in valid]
    return cols


def _latest_statement_col(df: pd.DataFrame):
    cols = _ordered_statement_columns(df)
    return cols[-1] if cols else None


def _format_period_label(col):
    if col is None:
        return None
    if isinstance(col, (pd.Timestamp, datetime)):
        return col.strftime("%Y-%m-%d")
    try:
        dt = pd.to_datetime(col, errors="coerce")
        if pd.notna(dt):
            return dt.strftime("%Y-%m-%d")
    except Exception:
        pass
    return str(col)


def _pick_statement_value(df: pd.DataFrame, keys, col):
    if df is None or df.empty or col is None:
        return None
    for key in keys:
        if key in df.index:
            return _safe_float(df.at[key, col])
    return None


def _format_compact_currency(value, currency="USD"):
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "--"
    try:
        value = float(value)
    except Exception:
        return str(value)
    abs_val = abs(value)
    if abs_val >= 1e12:
        return f"{value / 1e12:,.2f}T {currency}"
    if abs_val >= 1e9:
        return f"{value / 1e9:,.2f}B {currency}"
    if abs_val >= 1e6:
        return f"{value / 1e6:,.2f}M {currency}"
    if abs_val >= 1e3:
        return f"{value / 1e3:,.2f}K {currency}"
    return f"{value:,.0f} {currency}"


def _format_amount(value, scale=1e6, decimals=1):
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "--"
    try:
        value = float(value)
    except Exception:
        return str(value)
    return f"{value / scale:,.{decimals}f}"


def _format_percent(value, decimals=1):
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "--"
    try:
        value = float(value)
    except Exception:
        return str(value)
    return f"{value * 100:.{decimals}f}%"


def _parse_rate(raw, default):
    if raw is None or raw == "":
        return default
    try:
        value = float(str(raw).replace(",", "."))
    except Exception:
        return default
    if value > 1:
        value = value / 100.0
    return value


def _parse_int(raw, default, min_value=None, max_value=None):
    try:
        value = int(raw)
    except Exception:
        value = default
    if min_value is not None:
        value = max(min_value, value)
    if max_value is not None:
        value = min(max_value, value)
    return value


def obter_snapshot_preco(ticker):
    """Return current price, previous close and intraday change percentage."""
    try:
        ticker_obj = yf.Ticker(ticker)
    except Exception:
        return {"price": None, "previous_close": None, "change_pct": None}

    fast_info = None
    try:
        fast_info = ticker_obj.fast_info
    except Exception:
        fast_info = None

    price = _pick_float(fast_info, ("regularMarketPrice", "lastPrice", "last_price"))
    previous_close = _pick_float(
        fast_info,
        ("regularMarketPreviousClose", "previousClose", "previous_close"),
    )

    if price is None:
        try:
            hist = ticker_obj.history(period="1d", interval="1m")
            if not hist.empty:
                closes = hist["Close"].dropna()
                if not closes.empty:
                    price = _safe_float(closes.iloc[-1])
        except Exception:
            pass

    if previous_close is None:
        try:
            hist_daily = ticker_obj.history(period="5d", interval="1d", auto_adjust=False)
            if not hist_daily.empty:
                closes = hist_daily["Close"].dropna()
                if len(closes) >= 2:
                    previous_close = _safe_float(closes.iloc[-2])
                elif len(closes) == 1:
                    previous_close = _safe_float(closes.iloc[-1])
        except Exception:
            pass

    change_pct = None
    if price is not None and previous_close not in (None, 0):
        change_pct = ((price - previous_close) / previous_close) * 100

    return {"price": price, "previous_close": previous_close, "change_pct": change_pct}


def obter_preco_atual(ticker):
    snapshot = obter_snapshot_preco(ticker)
    return snapshot.get("price")


def baixar_dados(ticker, period="6mo", interval="1d"):
    df = yf.download(ticker, period=period, interval=interval, auto_adjust=False, progress=False)
    if df is None or df.empty:
        return pd.DataFrame()

    if isinstance(df.columns, pd.MultiIndex):
        df.columns = df.columns.get_level_values(0)

    cols = ["Open", "High", "Low", "Close", "Volume"]
    try:
        df = df[cols].copy()
    except KeyError:
        return pd.DataFrame()

    for col in cols:
        df[col] = pd.to_numeric(df[col], errors="coerce")

    df = df.dropna(how="any")
    return df

def calcular_indicadores(df):
    if df is None or df.empty:
        return pd.DataFrame()

    out = df.copy()

    numeric_cols = ["Open", "High", "Low", "Close", "Volume"]
    for col in numeric_cols:
        if col in out.columns:
            out[col] = pd.to_numeric(out[col], errors="coerce")

    out = out.dropna(subset=numeric_cols)
    if out.empty:
        return pd.DataFrame()

    def _empty_series(name):
        return pd.Series([float("nan")] * len(out), index=out.index, name=name)

    def safe_series(name, func):
        try:
            series = func()
            if series is None:
                return _empty_series(name)
            if not isinstance(series, pd.Series):
                series = pd.Series(series, index=out.index, name=name)
            else:
                series = series.reindex(out.index)
            return series
        except Exception:
            return _empty_series(name)

    out["EMA9"] = safe_series("EMA9", lambda: ta.trend.EMAIndicator(out["Close"], window=9).ema_indicator())
    out["EMA26"] = safe_series("EMA26", lambda: ta.trend.EMAIndicator(out["Close"], window=26).ema_indicator())

    out["MACD"] = safe_series("MACD", lambda: ta.trend.MACD(out["Close"]).macd())
    out["MACD_Signal"] = safe_series("MACD_Signal", lambda: ta.trend.MACD(out["Close"]).macd_signal())

    out["RSI"] = safe_series("RSI", lambda: ta.momentum.RSIIndicator(out["Close"], window=14).rsi())
    out["ADX"] = safe_series("ADX", lambda: ta.trend.ADXIndicator(out["High"], out["Low"], out["Close"], window=14).adx())

    out["BB_Upper"] = safe_series("BB_Upper", lambda: ta.volatility.BollingerBands(out["Close"], window=20).bollinger_hband())
    out["BB_Lower"] = safe_series("BB_Lower", lambda: ta.volatility.BollingerBands(out["Close"], window=20).bollinger_lband())

    out["ATR"] = safe_series("ATR", lambda: ta.volatility.AverageTrueRange(out["High"], out["Low"], out["Close"], window=14).average_true_range())
    out["OBV"] = safe_series("OBV", lambda: ta.volume.OnBalanceVolumeIndicator(out["Close"], out["Volume"]).on_balance_volume())
    out["CMF"] = safe_series("CMF", lambda: ta.volume.ChaikinMoneyFlowIndicator(out["High"], out["Low"], out["Close"], out["Volume"], window=20).chaikin_money_flow())

    return out.dropna(how="all")

def analisar(df, lang: str = DEFAULT_LANG):
    text = ANALYSIS_TEXT.get(lang, ANALYSIS_TEXT[DEFAULT_LANG])
    df_ok = df.dropna(subset=["Close", "EMA9", "EMA26"], how="any")
    if df_ok.empty:
        return {
            "close": None,
            "EMA9": None,
            "EMA26": None,
            "RSI": None,
            "ADX": None,
            "MACD": None,
            "MACD_Signal": None,
            "ATR": None,
            "CMF": None,
            "recomendacao": "NEUTRAL / WAIT",
            "recomendacao_label": text["rec_neutral"],
            "rationale": [text["insights_unavailable"]]
        }

    last_idx = df_ok.index[-1]

    def last_value(col):
        if col not in df.columns:
            return float("nan")
        series = df.loc[:last_idx, col].dropna()
        if series.empty:
            return float("nan")
        try:
            return float(series.iloc[-1])
        except Exception:
            return float("nan")

    def safe_round(value, digits):
        if value is None or pd.isna(value):
            return None
        try:
            return round(float(value), digits)
        except Exception:
            return None

    close = last_value("Close")
    ema9 = last_value("EMA9")
    ema26 = last_value("EMA26")
    macd = last_value("MACD")
    macd_sig = last_value("MACD_Signal")
    rsi = last_value("RSI")
    adx = last_value("ADX")
    atr = last_value("ATR")
    cmf = last_value("CMF")

    recomendacao = "NEUTRAL / WAIT"
    rationale = []
    summary = ""

    trend_up = (close > ema9 > ema26) if all(pd.notna([close, ema9, ema26])) else False
    trend_down = (close < ema9 < ema26) if all(pd.notna([close, ema9, ema26])) else False

    macd_now = (macd - macd_sig) if all(pd.notna([macd, macd_sig])) else float("nan")
    prev_macd = df_ok["MACD"].iloc[-2] if len(df_ok) > 1 and "MACD" in df_ok.columns else float("nan")
    prev_macd_sig = df_ok["MACD_Signal"].iloc[-2] if len(df_ok) > 1 and "MACD_Signal" in df_ok.columns else float("nan")
    if all(pd.notna([prev_macd, prev_macd_sig])):
        macd_prev = float(prev_macd) - float(prev_macd_sig)
    else:
        macd_prev = float("nan")
    macd_cross_down = (macd_prev >= 0 > macd_now) if all(pd.notna([macd_prev, macd_now])) else False

    adx_strong = adx >= 25 if pd.notna(adx) else False

    rsi_bull = (RISK_OPT_PARAMS["rsi_entry_min"] <= rsi <= RISK_OPT_PARAMS["rsi_entry_max"]) if pd.notna(rsi) else False
    rsi_hot = (rsi >= RISK_OPT_PARAMS["rsi_exit_max"]) if pd.notna(rsi) else False
    rsi_cold = (rsi <= 30) if pd.notna(rsi) else False

    bb_upper = last_value("BB_Upper")
    bb_lower = last_value("BB_Lower")
    bb_break_up = (close > bb_upper) if all(pd.notna([close, bb_upper])) else False
    bb_break_down = (close < bb_lower) if all(pd.notna([close, bb_lower])) else False

    if trend_up and (pd.notna(macd_now) and macd_now > 0) and rsi_bull and (pd.notna(adx) and adx >= RISK_OPT_PARAMS["adx_min"]):
        recomendacao = "ENTRY (buy)"
        summary = text["summary_entry"]
        rationale += [
            text["trend_positive"],
            text["macd_hist_pos"],
            text["rsi_between"].format(min=RISK_OPT_PARAMS["rsi_entry_min"], max=RISK_OPT_PARAMS["rsi_entry_max"]),
            text["adx_confirms"]
        ]
        if adx_strong:
            rationale.append(text["trend_strength"])
    elif trend_down or macd_cross_down or rsi_hot:
        recomendacao = "EXIT / AVOID"
        summary = text["summary_exit"]
        rationale += [
            text["trend_weakening"],
            text["rsi_take_profit"].format(max=RISK_OPT_PARAMS["rsi_exit_max"])
        ]
        if adx_strong:
            rationale.append(text["trend_strength"])
    else:
        if rsi_hot and bb_break_up:
            summary = text["summary_overbought"]
            rationale += [text["overbought_pullback"]]
        elif rsi_cold and bb_break_down:
            summary = text["summary_oversold"]
            rationale += [text["oversold_reversal"]]
        else:
            summary = text["summary_mixed"]
            if trend_up:
                rationale += [text["trend_up_wait"]]
            elif trend_down:
                rationale += [text["trend_down_wait"]]
            else:
                rationale += [text["sideways_wait"]]

        if not rationale:
            rationale.append(text["summary_none"])

    return {
        "close": safe_round(close, 2),
        "EMA9": safe_round(ema9, 2),
        "EMA26": safe_round(ema26, 2),
        "RSI": safe_round(rsi, 2),
        "ADX": safe_round(adx, 2),
        "MACD": safe_round(macd, 3),
        "MACD_Signal": safe_round(macd_sig, 3),
        "ATR": safe_round(atr, 3),
        "CMF": safe_round(cmf, 3),
        "recomendacao": recomendacao,
        "recomendacao_label": (
            text["rec_entry"] if "ENTRY" in recomendacao
            else text["rec_exit"] if "EXIT" in recomendacao
            else text["rec_neutral"]
        ),
        "rationale": ([summary] if summary else []) + rationale
    }

# =============================
# Business helpers
# =============================


def process_tickers(tickers, lang: str = DEFAULT_LANG):
    results = {}
    for ticker in tickers:
        try:
            df = baixar_dados(ticker)
            if df.empty:
                results[ticker] = {"erro": "No data available"}
                continue
            prev_close_change_pct = _calc_prev_close_change_pct(df)

            df = calcular_indicadores(df)
            if df.empty:
                results[ticker] = {"erro": "Indicators unavailable"}
                continue

            analisado = analisar(df, lang)
            snapshot = obter_snapshot_preco(ticker)
            live_price = snapshot.get("price")
            previous_close = snapshot.get("previous_close")
            change_pct = snapshot.get("change_pct")

            if live_price is not None:
                analisado["price"] = round(float(live_price), 2)
                analisado["price_source"] = "live"
            else:
                analisado["price"] = analisado.get("close")
                analisado["price_source"] = "previous_close" if analisado.get("close") is not None else None

            if previous_close is None:
                previous_close = analisado.get("close")
            analisado["previous_close"] = round(float(previous_close), 2) if previous_close is not None else None

            if change_pct is None and live_price is not None and previous_close not in (None, 0):
                change_pct = ((float(live_price) - float(previous_close)) / float(previous_close)) * 100
            change_pct_today = round(float(change_pct), 2) if change_pct is not None else None
            analisado["change_pct_today"] = change_pct_today
            analisado["change_pct_prev_close"] = (
                round(float(prev_close_change_pct), 2) if prev_close_change_pct is not None else None
            )
            analisado["change_pct"] = change_pct_today

            results[ticker] = analisado
        except Exception as exc:
            results[ticker] = {"erro": str(exc)}

    return results


def _calc_cagr(values):
    cleaned = [v for v in values if v is not None]
    if len(cleaned) < 2:
        return None
    first = cleaned[0]
    last = cleaned[-1]
    if first <= 0 or last <= 0:
        return None
    years = len(cleaned) - 1
    try:
        return (last / first) ** (1 / years) - 1
    except Exception:
        return None

def _calc_ev_simple(fcf_base, growth_rate, discount_rate, terminal_growth, years):
    if fcf_base is None or years <= 0:
        return None
    if discount_rate <= terminal_growth:
        return None
    projected_fcf = fcf_base
    pv_total = 0.0
    for year in range(1, years + 1):
        projected_fcf = projected_fcf * (1 + growth_rate)
        pv_total += projected_fcf / ((1 + discount_rate) ** year)
    terminal_value = (projected_fcf * (1 + terminal_growth)) / (discount_rate - terminal_growth)
    terminal_pv = terminal_value / ((1 + discount_rate) ** years)
    return pv_total + terminal_pv


def _growth_source_label(lang, source, years):
    years = years or 0
    if source == "fcf":
        if lang == "pt":
            return f"CAGR FCF ({years} anos)"
        if lang == "es":
            return f"CAGR FCF ({years} anos)"
        return f"FCF CAGR ({years}y)"
    if source == "revenue":
        if lang == "pt":
            return f"CAGR Receita ({years} anos)"
        if lang == "es":
            return f"CAGR Ingresos ({years} anos)"
        return f"Revenue CAGR ({years}y)"
    if lang == "pt":
        return "Padrao"
    if lang == "es":
        return "Estandar"
    return "Default"

def _local_text(lang, pt, en, es):
    if lang == "pt":
        return pt
    if lang == "es":
        return es
    return en


def build_dcf_context(ticker, lang: str = DEFAULT_LANG, params=None):
    lang = resolve_lang(lang)
    i18n = I18N.get(lang, I18N[DEFAULT_LANG])
    params = params or {}

    context = {
        "ticker": (ticker or "").strip().upper(),
        "lang": lang,
        "i18n": i18n,
        "erro": None,
        "currency": "USD",
        "company_name": None,
        "base_period": None,
        "base_snapshot": [],
        "base_note": None,
        "latest_metrics": [],
        "history_rows": [],
        "assumptions": [],
        "dcf_rows": [],
        "valuation_metrics": [],
        "pv_total_fmt": None,
        "terminal_value_fmt": None,
        "terminal_pv_fmt": None,
        "enterprise_value_fmt": None,
        "market_cap_fmt": None,
        "pv_sum_label": None,
        "terminal_formula": None,
        "interpretation_points": [],
        "sensitivity_notes": [],
        "conclusion_points": [],
        "sensitivity": None,
        "notes": [],
        "amount_unit": None,
    }
    context["company_name"] = context["ticker"] or "--"

    if not context["ticker"]:
        context["erro"] = i18n.get("dcf_missing_data")
        return context

    try:
        ticker_obj = yf.Ticker(context["ticker"])
    except Exception as exc:
        context["erro"] = str(exc)
        return context

    info = {}
    try:
        info = ticker_obj.info or {}
    except Exception:
        info = {}

    try:
        fast_info = ticker_obj.fast_info or {}
    except Exception:
        fast_info = {}

    currency = info.get("currency") or "USD"
    context["currency"] = currency
    company_name = info.get("longName") or info.get("shortName") or context["ticker"]
    context["company_name"] = company_name

    try:
        cashflow = ticker_obj.cashflow
    except Exception:
        cashflow = pd.DataFrame()
    try:
        financials = ticker_obj.financials
    except Exception:
        financials = pd.DataFrame()
    try:
        balance = ticker_obj.balance_sheet
    except Exception:
        balance = pd.DataFrame()

    latest_col = (
        _latest_statement_col(cashflow)
        or _latest_statement_col(financials)
        or _latest_statement_col(balance)
    )
    latest_period = _format_period_label(latest_col) or "--"
    context["base_period"] = latest_period

    operating_cf = _pick_statement_value(
        cashflow,
        ["Total Cash From Operating Activities", "Operating Cash Flow", "OperatingCashFlow"],
        latest_col,
    )
    capex = _pick_statement_value(
        cashflow,
        ["Capital Expenditures", "Capital Expenditure", "Capital Expenditures - Fixed Assets", "Purchase Of PPE"],
        latest_col,
    )
    fcf = _pick_statement_value(
        cashflow,
        ["Free Cash Flow", "Free Cashflow", "FreeCashFlow"],
        latest_col,
    )
    if fcf is None and operating_cf is not None and capex is not None:
        fcf = operating_cf + capex

    revenue = _pick_statement_value(
        financials,
        ["Total Revenue", "TotalRevenue"],
        latest_col,
    )
    net_income = _pick_statement_value(
        financials,
        ["Net Income", "Net Income Common Stockholders", "NetIncome"],
        latest_col,
    )
    ebitda = _pick_statement_value(
        financials,
        ["EBITDA", "Ebitda"],
        latest_col,
    )
    if ebitda is None:
        ebitda = _safe_float(info.get("ebitda"))

    cash = _pick_statement_value(
        balance,
        [
            "Cash And Cash Equivalents",
            "Cash",
            "Cash And Short Term Investments",
            "Cash And Cash Equivalents And Short Term Investments",
        ],
        latest_col,
    )
    debt = _pick_statement_value(
        balance,
        [
            "Total Debt",
            "Long Term Debt",
            "Short Long Term Debt",
            "Short Term Debt",
            "Long Term Debt And Capital Lease Obligation",
        ],
        latest_col,
    )

    shares = _safe_float(info.get("sharesOutstanding"))
    price = _pick_float(
        fast_info,
        ("lastPrice", "last_price", "regularMarketPrice", "regularMarketLastPrice"),
    )
    if price is None:
        price = _safe_float(info.get("currentPrice"))
    market_cap = _safe_float(info.get("marketCap"))
    if shares is None and market_cap is not None and price not in (None, 0):
        shares = market_cap / price

    context["latest_metrics"] = [
        (i18n.get("dcf_field_fiscal"), latest_period),
        (i18n.get("dcf_field_currency"), currency),
        (i18n.get("dcf_field_revenue"), _format_compact_currency(revenue, currency)),
        (i18n.get("dcf_field_operating_cf"), _format_compact_currency(operating_cf, currency)),
        (i18n.get("dcf_field_capex"), _format_compact_currency(capex, currency)),
        (i18n.get("dcf_field_fcf"), _format_compact_currency(fcf, currency)),
        (i18n.get("dcf_field_ebitda"), _format_compact_currency(ebitda, currency)),
        (i18n.get("dcf_field_net_income"), _format_compact_currency(net_income, currency)),
        (i18n.get("dcf_field_cash"), _format_compact_currency(cash, currency)),
        (i18n.get("dcf_field_debt"), _format_compact_currency(debt, currency)),
        (i18n.get("dcf_field_shares"), f"{shares:,.0f}" if shares else "--"),
    ]

    context["base_snapshot"] = [
        (i18n.get("dcf_field_revenue"), _format_compact_currency(revenue, currency)),
        (i18n.get("dcf_field_net_income"), _format_compact_currency(net_income, currency)),
        (i18n.get("dcf_field_fcf"), _format_compact_currency(fcf, currency)),
    ]

    base_note = None
    if fcf is not None and net_income is not None:
        if fcf < net_income:
            base_note = _local_text(
                lang,
                "FCF abaixo do lucro liquido; capex elevado pode explicar.",
                "FCF below net income; elevated capex can explain.",
                "FCF por debajo de la utilidad neta; capex elevado puede explicarlo.",
            )
        elif fcf > net_income:
            base_note = _local_text(
                lang,
                "FCF acima do lucro liquido; forte geracao de caixa.",
                "FCF above net income; strong cash generation.",
                "FCF por encima de la utilidad neta; fuerte generacion de caja.",
            )
        else:
            base_note = _local_text(
                lang,
                "FCF proximo do lucro liquido.",
                "FCF close to net income.",
                "FCF cercano a la utilidad neta.",
            )
    context["base_note"] = base_note

    history_source = cashflow if cashflow is not None and not cashflow.empty else financials
    history_cols = _ordered_statement_columns(history_source)
    history_rows = []
    for col in history_cols[-4:]:
        row_revenue = _pick_statement_value(financials, ["Total Revenue", "TotalRevenue"], col)
        row_fcf = _pick_statement_value(cashflow, ["Free Cash Flow", "Free Cashflow", "FreeCashFlow"], col)
        if row_fcf is None:
            row_operating = _pick_statement_value(
                cashflow,
                ["Total Cash From Operating Activities", "Operating Cash Flow", "OperatingCashFlow"],
                col,
            )
            row_capex = _pick_statement_value(
                cashflow,
                ["Capital Expenditures", "Capital Expenditure", "Capital Expenditures - Fixed Assets", "Purchase Of PPE"],
                col,
            )
            if row_operating is not None and row_capex is not None:
                row_fcf = row_operating + row_capex
        history_rows.append({
            "period": _format_period_label(col) or "--",
            "revenue": _format_compact_currency(row_revenue, currency),
            "fcf": _format_compact_currency(row_fcf, currency),
        })
    context["history_rows"] = history_rows

    fcf_series = []
    for col in _ordered_statement_columns(cashflow):
        value = _pick_statement_value(cashflow, ["Free Cash Flow", "Free Cashflow", "FreeCashFlow"], col)
        if value is None:
            op_cf = _pick_statement_value(
                cashflow,
                ["Total Cash From Operating Activities", "Operating Cash Flow", "OperatingCashFlow"],
                col,
            )
            capex_val = _pick_statement_value(
                cashflow,
                ["Capital Expenditures", "Capital Expenditure", "Capital Expenditures - Fixed Assets", "Purchase Of PPE"],
                col,
            )
            if op_cf is not None and capex_val is not None:
                value = op_cf + capex_val
        fcf_series.append(value)

    revenue_series = [
        _pick_statement_value(financials, ["Total Revenue", "TotalRevenue"], col)
        for col in _ordered_statement_columns(financials)
    ]

    growth_rate = None
    growth_source = "default"
    if fcf_series:
        growth_rate = _calc_cagr([v for v in fcf_series if v is not None])
        if growth_rate is not None:
            growth_source = "fcf"
    if growth_rate is None and revenue_series:
        growth_rate = _calc_cagr([v for v in revenue_series if v is not None])
        if growth_rate is not None:
            growth_source = "revenue"
    if growth_rate is None:
        growth_rate = 0.04

    growth_rate = max(-0.05, min(0.15, growth_rate))
    fcf_years = max(0, len([v for v in fcf_series if v is not None]) - 1)
    revenue_years = max(0, len([v for v in revenue_series if v is not None]) - 1)
    source_years = fcf_years if growth_source == "fcf" else revenue_years if growth_source == "revenue" else 0

    years = _parse_int(params.get("years"), 5, min_value=3, max_value=10)
    discount_rate = _parse_rate(params.get("discount"), 0.10)
    terminal_growth = _parse_rate(params.get("terminal"), 0.025)

    if terminal_growth >= discount_rate:
        terminal_growth = max(0.0, discount_rate - 0.01)
        context["notes"].append(i18n.get("dcf_note_terminal_adjust"))

    context["assumptions"] = [
        (i18n.get("dcf_field_years"), str(years)),
        (i18n.get("dcf_field_discount_rate"), _format_percent(discount_rate)),
        (i18n.get("dcf_field_terminal_growth"), _format_percent(terminal_growth)),
        (i18n.get("dcf_field_growth"), _format_percent(growth_rate)),
        (i18n.get("dcf_field_growth_source"), _growth_source_label(lang, growth_source, source_years)),
    ]

    if fcf is None:
        context["erro"] = i18n.get("dcf_missing_data")
        context["notes"].extend([i18n.get("dcf_note_source"), i18n.get("dcf_note_simplified")])
        return context

    if fcf <= 0:
        if lang == "pt":
            context["notes"].append("FCF negativo no ultimo ano; o DCF fica mais sensivel.")
        elif lang == "es":
            context["notes"].append("FCF negativo en el ultimo ano; el DCF es mas sensible.")
        else:
            context["notes"].append("Negative FCF in the latest year; DCF becomes more sensitive.")

    scale = 1e6
    context["amount_unit"] = f"{currency} mm"

    dcf_rows = []
    pv_total = 0.0
    projected_fcf = fcf
    for year in range(1, years + 1):
        projected_fcf = projected_fcf * (1 + growth_rate)
        discount_factor = 1 / ((1 + discount_rate) ** year)
        pv = projected_fcf * discount_factor
        pv_total += pv
        dcf_rows.append({
            "year": year,
            "fcf": _format_amount(projected_fcf, scale=scale, decimals=1),
            "discount": f"{discount_factor:.3f}",
            "pv": _format_amount(pv, scale=scale, decimals=1),
        })
    context["dcf_rows"] = dcf_rows
    fcf_last = projected_fcf

    terminal_value = (fcf_last * (1 + terminal_growth)) / (discount_rate - terminal_growth)
    terminal_pv = terminal_value / ((1 + discount_rate) ** years)
    enterprise_value = pv_total + terminal_pv

    net_cash = 0.0
    if cash is not None:
        net_cash += cash
    if debt is not None:
        net_cash -= debt

    equity_value = enterprise_value + net_cash
    intrinsic_value = equity_value / shares if shares else None

    context["valuation_metrics"] = [
        (i18n.get("dcf_terminal_value"), _format_compact_currency(terminal_value, currency)),
        (i18n.get("dcf_terminal_pv"), _format_compact_currency(terminal_pv, currency)),
        (i18n.get("dcf_summary_enterprise"), _format_compact_currency(enterprise_value, currency)),
        (i18n.get("dcf_summary_equity"), _format_compact_currency(equity_value, currency)),
        (i18n.get("dcf_summary_intrinsic"), f"{intrinsic_value:,.2f} {currency}" if intrinsic_value else "--"),
        (i18n.get("dcf_summary_price"), f"{price:,.2f} {currency}" if price else "--"),
        (i18n.get("dcf_summary_upside"), _format_percent((intrinsic_value - price) / price) if intrinsic_value and price else "--"),
    ]

    context["pv_total_fmt"] = _format_compact_currency(pv_total, currency)
    context["terminal_value_fmt"] = _format_compact_currency(terminal_value, currency)
    context["terminal_pv_fmt"] = _format_compact_currency(terminal_pv, currency)
    context["enterprise_value_fmt"] = _format_compact_currency(enterprise_value, currency)
    context["market_cap_fmt"] = _format_compact_currency(market_cap, currency) if market_cap is not None else None
    context["pv_sum_label"] = (i18n.get("dcf_pv_sum") or "PV sum").format(years=years)

    terminal_formula = None
    if fcf_last is not None:
        fcf_last_fmt = _format_amount(fcf_last, scale=scale, decimals=1)
        terminal_formula = _local_text(
            lang,
            f"VT = FCF_{years} ({fcf_last_fmt} {context['amount_unit']}) x (1 + {terminal_growth * 100:.1f}%) / ({discount_rate * 100:.1f}% - {terminal_growth * 100:.1f}%) = {context['terminal_value_fmt']}",
            f"TV = FCF_{years} ({fcf_last_fmt} {context['amount_unit']}) x (1 + {terminal_growth * 100:.1f}%) / ({discount_rate * 100:.1f}% - {terminal_growth * 100:.1f}%) = {context['terminal_value_fmt']}",
            f"VT = FCF_{years} ({fcf_last_fmt} {context['amount_unit']}) x (1 + {terminal_growth * 100:.1f}%) / ({discount_rate * 100:.1f}% - {terminal_growth * 100:.1f}%) = {context['terminal_value_fmt']}",
        )
    context["terminal_formula"] = terminal_formula

    interpretation_points = []
    conclusion_points = []
    if market_cap is not None and market_cap > 0:
        diff_pct = (enterprise_value - market_cap) / market_cap
        market_cap_text = context["market_cap_fmt"]
        if diff_pct <= -0.1:
            interpretation_points.append(
                _local_text(
                    lang,
                    f"O valuation estimado esta abaixo do market cap atual ({market_cap_text}).",
                    f"The estimated valuation is below the current market cap ({market_cap_text}).",
                    f"El valuation estimado esta por debajo del market cap actual ({market_cap_text}).",
                )
            )
            interpretation_points.append(
                _local_text(
                    lang,
                    "O mercado pode estar precificando crescimento maior, FCF estrutural mais alto ou expansao de margens.",
                    "The market may be pricing higher growth, structurally higher FCF, or margin expansion.",
                    "El mercado puede estar descontando mayor crecimiento, FCF estructural mas alto o expansion de margenes.",
                )
            )
            conclusion_points.append(
                _local_text(
                    lang,
                    "Mercado precifica um cenario mais agressivo do que o conservador do modelo.",
                    "The market is pricing a more aggressive scenario than this conservative model.",
                    "El mercado descuenta un escenario mas agresivo que el conservador del modelo.",
                )
            )
        elif diff_pct >= 0.1:
            interpretation_points.append(
                _local_text(
                    lang,
                    f"O valuation estimado esta acima do market cap atual ({market_cap_text}).",
                    f"The estimated valuation is above the current market cap ({market_cap_text}).",
                    f"El valuation estimado esta por encima del market cap actual ({market_cap_text}).",
                )
            )
            interpretation_points.append(
                _local_text(
                    lang,
                    "O mercado pode estar mais conservador quanto a crescimento ou margens.",
                    "The market may be more conservative on growth or margins.",
                    "El mercado puede ser mas conservador en crecimiento o margenes.",
                )
            )
            conclusion_points.append(
                _local_text(
                    lang,
                    "Modelo sugere upside vs. market cap atual, mas depende das premissas.",
                    "The model suggests upside vs. the current market cap, but it depends on the assumptions.",
                    "El modelo sugiere potencial alcista vs. el market cap actual, pero depende de las premisas.",
                )
            )
        else:
            interpretation_points.append(
                _local_text(
                    lang,
                    f"O valuation estimado esta proximo do market cap atual ({market_cap_text}).",
                    f"The estimated valuation is close to the current market cap ({market_cap_text}).",
                    f"El valuation estimado esta cerca del market cap actual ({market_cap_text}).",
                )
            )
            conclusion_points.append(
                _local_text(
                    lang,
                    "Premissas do modelo parecem alinhadas ao consenso de mercado.",
                    "Model assumptions appear aligned with market consensus.",
                    "Las premisas del modelo parecen alineadas con el consenso de mercado.",
                )
            )
    else:
        interpretation_points.append(
            _local_text(
                lang,
                "Market cap indisponivel; use o EV como referencia comparativa.",
                "Market cap unavailable; use EV as a comparative reference.",
                "Market cap no disponible; use el EV como referencia comparativa.",
            )
        )

    context["interpretation_points"] = interpretation_points

    sensitivity_notes = []
    if fcf is not None and fcf > 0:
        growth_hi = min(growth_rate + 0.05, 0.25)
        if growth_hi > growth_rate:
            ev_hi = _calc_ev_simple(fcf, growth_hi, discount_rate, terminal_growth, years)
            if ev_hi is not None:
                sensitivity_notes.append(
                    _local_text(
                        lang,
                        f"Se crescimento FCF for {growth_hi * 100:.0f}% por {years} anos, EV ~ {_format_compact_currency(ev_hi, currency)}.",
                        f"If FCF growth is {growth_hi * 100:.0f}% for {years} years, EV ~ {_format_compact_currency(ev_hi, currency)}.",
                        f"Si el crecimiento de FCF es {growth_hi * 100:.0f}% por {years} anos, EV ~ {_format_compact_currency(ev_hi, currency)}.",
                    )
                )
        fcf_hi = fcf * 1.5
        ev_fcf_hi = _calc_ev_simple(fcf_hi, growth_rate, discount_rate, terminal_growth, years)
        if ev_fcf_hi is not None:
            sensitivity_notes.append(
                _local_text(
                    lang,
                    f"Se FCF base subir para {_format_compact_currency(fcf_hi, currency)}, EV ~ {_format_compact_currency(ev_fcf_hi, currency)}.",
                    f"If base FCF rises to {_format_compact_currency(fcf_hi, currency)}, EV ~ {_format_compact_currency(ev_fcf_hi, currency)}.",
                    f"Si el FCF base sube a {_format_compact_currency(fcf_hi, currency)}, EV ~ {_format_compact_currency(ev_fcf_hi, currency)}.",
                )
            )
    if sensitivity_notes:
        sensitivity_notes.append(
            _local_text(
                lang,
                "Veja a tabela de sensibilidade abaixo para variacoes de WACC e crescimento terminal.",
                "See the sensitivity table below for WACC and terminal growth variations.",
                "Vea la tabla de sensibilidad abajo para variaciones de WACC y crecimiento terminal.",
            )
        )
    context["sensitivity_notes"] = sensitivity_notes
    context["conclusion_points"] = conclusion_points

    if shares and price:
        discount_rates = [
            max(0.01, discount_rate - 0.01),
            discount_rate,
            discount_rate + 0.01,
        ]
        terminal_rates = [
            max(0.0, terminal_growth - 0.005),
            terminal_growth,
            terminal_growth + 0.005,
        ]
        sensitivity = []
        for tr in terminal_rates:
            row = []
            for dr in discount_rates:
                if tr >= dr:
                    row.append("--")
                    continue
                terminal_val = (projected_fcf * (1 + tr)) / (dr - tr)
                terminal_pv_val = terminal_val / ((1 + dr) ** years)
                ev_val = 0.0
                for y in range(1, years + 1):
                    fcf_y = fcf * ((1 + growth_rate) ** y)
                    ev_val += fcf_y / ((1 + dr) ** y)
                ev_val += terminal_pv_val
                equity_val = ev_val + net_cash
                per_share_val = equity_val / shares
                row.append(f"{per_share_val:,.2f}")
            sensitivity.append(row)
        sensitivity_rows = []
        for idx, tr in enumerate(terminal_rates):
            sensitivity_rows.append({
                "terminal": f"{tr * 100:.1f}%",
                "cells": sensitivity[idx],
            })
        context["sensitivity"] = {
            "discount_rates": [f"{dr * 100:.1f}%" for dr in discount_rates],
            "rows": sensitivity_rows,
        }

    context["notes"].extend([i18n.get("dcf_note_source"), i18n.get("dcf_note_simplified")])

    return context

def generate_indicator_doc(symbols, resultados=None, lang: str = DEFAULT_LANG):
    if resultados is None:
        resultados = process_tickers(symbols, lang)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    base_style = doc.styles["Normal"]
    base_font = base_style.font
    base_font.name = "Segoe UI"
    base_font.size = Pt(10)
    base_font.color.rgb = RGBColor(232, 234, 237)

    COLOR_TEXT = RGBColor(232, 234, 237)
    COLOR_MUTED = RGBColor(167, 176, 190)
    COLOR_ACCENT = RGBColor(93, 169, 255)
    COLOR_OK = RGBColor(0, 194, 138)
    COLOR_DANGER = RGBColor(255, 107, 107)
    COLOR_WARN = RGBColor(255, 176, 32)

    BG_HEADER = "171A21"
    BG_ROWS = ["181D26", "141821"]
    BG_PILL_BUY = "123329"
    BG_PILL_SELL = "341F20"
    BG_PILL_NEUTRAL = "2F2615"
    BG_PILL_DISABLED = "232730"

    def set_table_borders(table, color_hex="262B36"):
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.append(tbl_pr)
        tbl_borders = tbl_pr.find(qn("w:tblBorders"))
        if tbl_borders is None:
            tbl_borders = OxmlElement("w:tblBorders")
            tbl_pr.append(tbl_borders)
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = tbl_borders.find(qn(f"w:{border_name}"))
            if border is None:
                border = OxmlElement(f"w:{border_name}")
                tbl_borders.append(border)
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), color_hex)

    def set_cell_background(cell, hex_color):
        tc_pr = cell._tc.get_or_add_tcPr()
        existing = tc_pr.findall(qn("w:shd"))
        for shd in existing:
            tc_pr.remove(shd)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), hex_color)
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:val"), "clear")
        tc_pr.append(shading)

    def set_cell_text(cell, text, color=COLOR_TEXT, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = align
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.color.rgb = color
        run.font.size = Pt(10)

    def fmt_value(value, decimals=2):
        if value is None or (isinstance(value, (float, int)) and (pd.isna(value))):
            return "N/A"
        try:
            return f"{float(value):.{decimals}f}"
        except (TypeError, ValueError):
            return str(value)

    if LOGO_PATH and os.path.exists(LOGO_PATH):
        doc.add_picture(LOGO_PATH, width=Inches(2.8))
        logo_paragraph = doc.paragraphs[-1]
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_paragraph.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title_run = title.add_run("Indicator List & Recommendations")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = COLOR_ACCENT

    meta = doc.add_paragraph(f"Generated on {datetime.utcnow().isoformat(timespec='seconds')}Z")
    meta_run = meta.runs[0]
    meta_run.font.color.rgb = COLOR_MUTED
    meta_run.font.size = Pt(10)

    doc.add_paragraph()

    headers = [
        "Ticker",
        "Price",
        "RSI",
        "MACD / Signal",
        "ADX",
        "CMF",
        "Recommendation",
        "Summary / Rationale",
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)

    header_row = table.rows[0]
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        set_cell_background(cell, BG_HEADER)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != len(headers) - 1 else WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = COLOR_ACCENT
        run.font.size = Pt(10)

    column_widths = [
        Inches(1.0),
        Inches(0.9),
        Inches(0.8),
        Inches(1.5),
        Inches(0.8),
        Inches(0.8),
        Inches(1.4),
        Inches(3.0),
    ]

    for row_idx, ticker in enumerate(symbols):
        data = resultados.get(ticker, {})
        erro = data.get("erro")
        row = table.add_row()
        row_background = BG_ROWS[row_idx % len(BG_ROWS)]
        for cell in row.cells:
            set_cell_background(cell, row_background)

        ticker_color = COLOR_MUTED if erro else COLOR_ACCENT
        set_cell_text(row.cells[0], ticker, color=ticker_color, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

        def fill_metric(cell, value, decimals=2):
            if erro:
                set_cell_text(cell, "--", color=COLOR_MUTED)
            else:
                set_cell_text(cell, fmt_value(value, decimals), color=COLOR_TEXT)

        price_display = data.get("price")
        if price_display is None:
            price_display = data.get("close")
        if erro:
            set_cell_text(row.cells[1], "--", color=COLOR_MUTED)
        else:
            price_text = fmt_value(price_display, 2)
            if data.get("price_source") == "live":
                price_text = f"{price_text} (live)"
            elif data.get("price_source") == "previous_close":
                price_text = f"{price_text} (prev close)"
            if price_text == "N/A":
                set_cell_text(row.cells[1], price_text, color=COLOR_MUTED)
            else:
                set_cell_text(row.cells[1], price_text, color=COLOR_TEXT)
        fill_metric(row.cells[2], data.get("RSI"), 2)
        if erro:
            set_cell_text(row.cells[3], "-- / --", color=COLOR_MUTED)
        else:
            macd = fmt_value(data.get("MACD"), 3)
            macd_sig = fmt_value(data.get("MACD_Signal"), 3)
            set_cell_text(row.cells[3], f"{macd} / {macd_sig}", color=COLOR_TEXT)
        fill_metric(row.cells[4], data.get("ADX"), 2)
        fill_metric(row.cells[5], data.get("CMF"), 3)

        rec_cell = row.cells[6]
        rec_cell.text = ""
        if erro:
            set_cell_background(rec_cell, BG_PILL_DISABLED)
            paragraph = rec_cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("Unavailable")
            run.font.bold = True
            run.font.color.rgb = COLOR_MUTED
            run.font.size = Pt(10)
        else:
            recomendacao_code = data.get("recomendacao", "N/A")
            recomendacao_label = data.get("recomendacao_label") or recomendacao_code
            if "ENTRY" in (recomendacao_code or ""):
                pill_color = BG_PILL_BUY
                text_color = COLOR_OK
            elif "EXIT" in (recomendacao_code or ""):
                pill_color = BG_PILL_SELL
                text_color = COLOR_DANGER
            else:
                pill_color = BG_PILL_NEUTRAL
                text_color = COLOR_WARN
            set_cell_background(rec_cell, pill_color)
            paragraph = rec_cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(recomendacao_label)
            run.font.bold = True
            run.font.color.rgb = text_color
            run.font.size = Pt(10)

        summary_cell = row.cells[7]
        summary_cell.text = ""
        if erro:
            paragraph = summary_cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(erro)
            run.font.color.rgb = COLOR_MUTED
            run.font.size = Pt(10)
        else:
            rationale = data.get("rationale") or []
            if not rationale:
                paragraph = summary_cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.add_run("No rationale provided.")
                run.font.color.rgb = COLOR_MUTED
                run.font.size = Pt(10)
            else:
                first = True
                for item in rationale:
                    if first:
                        paragraph = summary_cell.paragraphs[0]
                        paragraph.style = "List Bullet"
                        first = False
                    else:
                        paragraph = summary_cell.add_paragraph(style="List Bullet")
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = paragraph.add_run(item)
                    run.font.color.rgb = COLOR_TEXT
                    run.font.size = Pt(10)

    for row in table.rows:
        for idx, width in enumerate(column_widths):
            row.cells[idx].width = width

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def send_indicator_report_email(ticker_string=None, lang: str = DEFAULT_LANG):
    recipients = [addr.strip() for addr in REPORT_EMAIL_TO.split(",") if addr.strip()]
    if not recipients:
        raise RuntimeError("REPORT_EMAIL_TO is not configured or empty.")
    if not REPORT_SMTP_HOST:
        raise RuntimeError("REPORT_SMTP_HOST is not configured.")

    ticker_string = normalize_ticker_string(ticker_string or REPORT_EMAIL_TICKERS)
    symbols = split_tickers(ticker_string)
    if not symbols:
        raise ValueError("No tickers provided for email report.")

    resultados = process_tickers(symbols, lang)
    report_bytes = generate_indicator_doc(symbols, resultados, lang)
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filename = f"indicator_list_{timestamp}.docx"

    msg = EmailMessage()
    msg["Subject"] = REPORT_EMAIL_SUBJECT
    msg["From"] = REPORT_EMAIL_FROM or REPORT_SMTP_USER or "reports@primesphere.local"
    msg["To"] = ", ".join(recipients)
    msg["Date"] = formatdate(localtime=False)

    body_text = REPORT_EMAIL_BODY or "Segue o relatÃ³rio diÃ¡rio de indicadores da PrimeSphere."
    msg.set_content(f"{body_text}\n\nTickers: {', '.join(symbols)}")
    if REPORT_EMAIL_BODY_HTML:
        msg.add_alternative(REPORT_EMAIL_BODY_HTML, subtype="html")

    msg.add_attachment(
        report_bytes,
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

    with smtplib.SMTP(REPORT_SMTP_HOST, REPORT_SMTP_PORT, timeout=30) as smtp:
        smtp.ehlo()
        if REPORT_SMTP_STARTTLS:
            context = ssl.create_default_context()
            smtp.starttls(context=context)
            smtp.ehlo()
        if REPORT_SMTP_USER:
            smtp.login(REPORT_SMTP_USER, REPORT_SMTP_PASSWORD or "")
        smtp.send_message(msg)

    return {
        "recipients": recipients,
        "symbols": symbols,
        "filename": filename,
    }

# =============================
# Rotas
# =============================
@app.route("/", methods=["GET","POST"])
def index():
    resultados = {}
    erro = None
    tickers = normalize_ticker_string(DEFAULT_TICKERS)
    lang = resolve_lang(request.values.get("lang"))

    if request.method == "POST":
        tickers = normalize_ticker_string(request.form.get("tickers", ""))
        symbols = split_tickers(tickers)

        if not symbols:
            erro = "Please enter at least one ticker symbol."
        else:
            resultados = process_tickers(symbols, lang)
    else:
        symbols = split_tickers(tickers)
        if symbols:
            resultados = process_tickers(symbols, lang)

    return render_template("lista.html",
                           resultados=resultados,
                           tickers=tickers,
                           erro=erro,
                           graph_base=GRAPH_APP_URL,
                           lang=lang,
                           i18n=I18N.get(lang, I18N[DEFAULT_LANG]))


@app.route("/dcf/<ticker>")
def dcf_view(ticker):
    lang = resolve_lang(request.args.get("lang"))
    params = {
        "years": request.args.get("years"),
        "discount": request.args.get("discount"),
        "terminal": request.args.get("terminal"),
    }
    context = build_dcf_context(ticker, lang, params)
    return render_template("dcf.html", **context)

@app.route("/api/recommendations")
def api_recommendations():
    tickers_raw = normalize_ticker_string(request.args.get("tickers", ""))
    lang = resolve_lang(request.args.get("lang"))
    symbols = split_tickers(tickers_raw)
    if not symbols:
        resp = jsonify({"error": "Please provide at least one ticker symbol."})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp, 400

    resultados = process_tickers(symbols, lang)
    data = {}
    for tk, res in resultados.items():
        if res.get("erro"):
            data[tk] = {"erro": res.get("erro")}
        else:
            data[tk] = {
                "recomendacao": res.get("recomendacao"),
                "recomendacao_label": res.get("recomendacao_label"),
            }

    payload = {
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "count": len(data),
        "data": data,
    }
    resp = jsonify(payload)
    resp.headers["Access-Control-Allow-Origin"] = "*"
    return resp

@app.route("/api/indicadores")
def api_indicadores():
    tickers_raw = normalize_ticker_string(request.args.get("tickers", ""))
    lang = resolve_lang(request.args.get("lang"))
    symbols = split_tickers(tickers_raw)
    if not symbols:
        resp = jsonify({"error": "Please provide at least one ticker symbol."})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp, 400

    resultados = process_tickers(symbols, lang)
    payload = {
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "count": len(resultados),
        "data": resultados,
    }
    resp = jsonify(payload)
    resp.headers["Access-Control-Allow-Origin"] = "*"
    return resp

@app.route("/logo")
def logo_image():
    logo_path = get_logo_path()
    if not logo_path:
        abort(404, description="Logo file not found.")
    logo_stream = get_logo_stream_no_background(logo_path)
    if logo_stream:
        return send_file(logo_stream, mimetype="image/png")
    return send_file(logo_path)

@app.route("/export/<fmt>")
def export_results(fmt):
    fmt = (fmt or "").lower()
    if fmt != "docx":
        abort(400, description="Unsupported export format.")

    tickers_raw = normalize_ticker_string(request.args.get("tickers", ""))
    lang = resolve_lang(request.args.get("lang"))
    symbols = split_tickers(tickers_raw)
    if not symbols:
        abort(400, description="Please provide at least one ticker symbol.")

    resultados = process_tickers(symbols, lang)

    report_bytes = generate_indicator_doc(symbols, resultados, lang)
    buffer = io.BytesIO(report_bytes)
    buffer.seek(0)
    filename = f"indicator_list_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"

    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == "__main__":
    import argparse
    import sys

    parser = argparse.ArgumentParser(description="PrimeSphere indicator list application.")
    parser.add_argument("--send-email", action="store_true", help="Generate and send the indicator report by e-mail.")
    parser.add_argument("--tickers", help="Comma-separated ticker symbols to override the default list.")
    parser.add_argument("--host", default="127.0.0.1", help="Host/IP for the Flask server (default: 127.0.0.1).")
    parser.add_argument("--port", type=int, default=5001, help="Port for the Flask server (default: 5001).")
    parser.add_argument("--debug", dest="debug", action="store_true", help="Run Flask with debug mode.")
    parser.add_argument("--no-debug", dest="debug", action="store_false", help="Run Flask without debug mode.")
    parser.set_defaults(debug=True)

    args = parser.parse_args()

    if args.send_email:
        try:
            info = send_indicator_report_email(args.tickers)
        except Exception as exc:
            print(f"[email] Failed to send report: {exc}", file=sys.stderr)
            sys.exit(1)
        print(f"[email] Report sent to {', '.join(info['recipients'])} ({info['filename']}).")
        print(f"[email] Symbols: {', '.join(info['symbols'])}")
        sys.exit(0)

    app.run(host=args.host, port=args.port, debug=args.debug)
